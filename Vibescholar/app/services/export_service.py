import io
import docx
from fpdf import FPDF
from typing import List
from app.models.reference import ProjectReference

class ExportService:
    @staticmethod
    def export_markdown(title: str, content: str, references: List[ProjectReference]) -> str:
        """
        Exports the document and its references as a unified Markdown string.
        """
        markdown = f"{content}\n\n"
        if references:
            markdown += "## Referências Bibliográficas\n\n"
            for ref in references:
                markdown += f"- {ref.authors}. **{ref.title}**. *{ref.journal or 'N/A'}*, {ref.year or 'N/A'}. DOI: {ref.doi or 'N/A'}.\n"
        return markdown

    @staticmethod
    def export_docx(title: str, content: str, references: List[ProjectReference]) -> bytes:
        """
        Generates a DOCX file containing the document content and bibliography.
        """
        doc = docx.Document()

        # Content parsing
        for block in content.split("\n\n"):
            block_stripped = block.strip()
            if not block_stripped:
                continue
            
            if block_stripped.startswith("# "):
                doc.add_heading(block_stripped[2:], level=1)
            elif block_stripped.startswith("## "):
                doc.add_heading(block_stripped[3:], level=2)
            elif block_stripped.startswith("### "):
                doc.add_heading(block_stripped[4:], level=3)
            else:
                doc.add_paragraph(block_stripped)

        # References Section
        if references:
            doc.add_heading("Referências Bibliográficas", level=1)
            for ref in references:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{ref.authors}. ")
                p.add_run(f"{ref.title}. ").bold = True
                p.add_run(f"{ref.journal or 'N/A'}, {ref.year or 'N/A'}. ")
                if ref.doi:
                    p.add_run(f"DOI: {ref.doi}")

        file_stream = io.BytesIO()
        doc.save(file_stream)
        return file_stream.getvalue()

    @staticmethod
    def export_pdf(title: str, content: str, references: List[ProjectReference]) -> bytes:
        """
        Generates a PDF using fpdf2.
        """
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)

        # Render Body Paragraphs
        pdf.set_font("Helvetica", size=11)
        for block in content.split("\n\n"):
            block_stripped = block.strip()
            if not block_stripped:
                continue
            
            if block_stripped.startswith("# "):
                pdf.ln(4)
                pdf.set_font("Helvetica", style="B", size=14)
                pdf.multi_cell(0, 8, txt=block_stripped[2:])
                pdf.set_font("Helvetica", size=11)
                pdf.ln(2)
            elif block_stripped.startswith("## "):
                pdf.ln(3)
                pdf.set_font("Helvetica", style="B", size=12)
                pdf.multi_cell(0, 8, txt=block_stripped[3:])
                pdf.set_font("Helvetica", size=11)
                pdf.ln(2)
            else:
                # Replace unsupported characters for standard latin-1 encoding
                text_clean = block_stripped.encode("latin-1", errors="replace").decode("latin-1")
                pdf.multi_cell(0, 6, txt=text_clean)
                pdf.ln(4)

        # Render References
        if references:
            pdf.ln(6)
            pdf.set_font("Helvetica", style="B", size=13)
            pdf.cell(0, 10, txt="Referencias Bibliograficas", ln=1)
            pdf.ln(2)
            pdf.set_font("Helvetica", size=10)
            
            for ref in references:
                ref_text = f"{ref.authors}. {ref.title}. {ref.journal or 'N/A'}, {ref.year or 'N/A'}."
                if ref.doi:
                    ref_text += f" DOI: {ref.doi}"
                clean_ref = ref_text.encode("latin-1", errors="replace").decode("latin-1")
                pdf.multi_cell(0, 6, txt=f"- {clean_ref}")
                pdf.ln(1)

        # fpdf2 may return bytearray depending on version; normalize for StreamingResponse.
        return bytes(pdf.output())

    @staticmethod
    def export_bibtex(references: List[ProjectReference]) -> str:
        """
        Formats references in BibTeX notation.
        """
        bib_entries = []
        for i, ref in enumerate(references):
            # Generate a key using first author lastname and year
            author_last = ref.authors.split(",")[0].split()[0].lower()
            year_key = ref.year if ref.year else "nodate"
            cite_key = f"{author_last}{year_key}_{ref.id}"
            
            entry = (
                f"@article{{{cite_key},\n"
                f"  author  = {{{ref.authors}}},\n"
                f"  title   = {{{ref.title}}},\n"
                f"  journal = {{{ref.journal or 'N/A'}}},\n"
                f"  year    = {{{ref.year or '0'}}}"
            )
            if ref.doi:
                entry += f",\n  doi     = {{{ref.doi}}}"
            entry += "\n}"
            bib_entries.append(entry)
            
        return "\n\n".join(bib_entries)

    @staticmethod
    def export_apa(references: List[ProjectReference]) -> str:
        """
        Formats references in APA style guidelines.
        """
        lines = []
        for ref in references:
            line = f"{ref.authors} ({ref.year or 'n.d.'}). {ref.title}."
            if ref.journal:
                line += f" {ref.journal}."
            if ref.doi:
                line += f" https://doi.org/{ref.doi}"
            lines.append(line)
        return "\n\n".join(lines)

    @staticmethod
    def export_abnt(references: List[ProjectReference]) -> str:
        """
        Formats references according to NBR 6023 (ABNT).
        """
        lines = []
        for ref in references:
            # Simple capitalized authors format representation
            line = f"{ref.authors.upper()}. **{ref.title}**."
            if ref.journal:
                line += f" {ref.journal},"
            if ref.year:
                line += f" {ref.year}."
            if ref.doi:
                line += f" Disponível em: <https://doi.org/{ref.doi}>."
            lines.append(line)
        return "\n\n".join(lines)
